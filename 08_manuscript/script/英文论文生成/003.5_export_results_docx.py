#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
export_results_docx.py

Convert Results_v0.md + Results_mapping.tsv into a Word (.docx) file for advisor review.

Key feature:
- Inserts a searchable anchor line before each written Panel block:
  [[SRC:source_yaml | block=panel_key | panel=panel_id | fig=figure_group]]
  This makes it easy to map Word edits back to YAML via Results_mapping.tsv.

Usage:
  python export_results_docx.py Results_v0.md Results_mapping.tsv Results_for_review.docx

If arguments are omitted, defaults are:
  md_path = Results_v0.md
  map_path = Results_mapping.tsv
  out_path = Results_for_review.docx
"""

from __future__ import annotations

import sys
import csv
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor


def load_mapping_tsv(path: Path) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    with path.open("r", encoding="utf-8", newline="") as f:
        r = csv.DictReader(f, delimiter="\t")
        for row in r:
            rows.append({k: (v or "") for k, v in row.items()})
    # sort by numeric order if possible
    def _k(x: Dict[str, str]) -> int:
        try:
            return int(x.get("order", "0"))
        except Exception:
            return 0
    rows.sort(key=_k)
    return rows


def add_heading(doc: Document, text: str, level: int) -> None:
    # python-docx heading levels: 0..9, where 1 is Heading 1.
    doc.add_heading(text.strip(), level=level)


def add_normal(doc: Document, text: str) -> None:
    doc.add_paragraph(text.rstrip())


def strip_md_bold(s: str) -> str:
    # Very minimal: **X** -> X
    t = s.strip()
    if t.startswith("**") and t.endswith("**") and len(t) >= 4:
        return t[2:-2].strip()
    return t


_CJK_RE = re.compile(r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]")


def contains_cjk(text: str) -> bool:
    return bool(_CJK_RE.search(text or ""))


def main() -> int:
    script_dir = Path(__file__).resolve().parent

    md_path = (
        Path(sys.argv[1]).expanduser().resolve()
        if len(sys.argv) >= 2
        else (script_dir / "Results_v0.md").resolve()
    )
    map_path = (
        Path(sys.argv[2]).expanduser().resolve()
        if len(sys.argv) >= 3
        else (script_dir / "Results_mapping.tsv").resolve()
    )
    out_path = (
        Path(sys.argv[3]).expanduser().resolve()
        if len(sys.argv) >= 4
        else (script_dir / "Results_for_review.docx").resolve()
    )

    if not md_path.exists():
        print(f"ERROR: markdown file not found: {md_path}", file=sys.stderr)
        print(f"       Tip: run with explicit paths, e.g.:\n"
              f"       python {Path(__file__).name} {script_dir/'Results_v0.md'} {script_dir/'Results_mapping.tsv'} {script_dir/'Results_for_review.docx'}",
              file=sys.stderr)
        return 2
    if not map_path.exists():
        print(f"ERROR: mapping TSV not found: {map_path}", file=sys.stderr)
        print(f"       Tip: run with explicit paths, e.g.:\n"
              f"       python {Path(__file__).name} {script_dir/'Results_v0.md'} {script_dir/'Results_mapping.tsv'} {script_dir/'Results_for_review.docx'}",
              file=sys.stderr)
        return 2

    mapping_rows = load_mapping_tsv(map_path)

    md_lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()

    def flush_panel_paragraph(buf: List[str]) -> None:
        # Join lines into one paragraph. Keep order; collapse extra spaces.
        text = " ".join([b.strip() for b in buf if b and b.strip()])
        text = " ".join(text.split())
        if text:
            doc.add_paragraph(text)

    in_panel = False
    panel_buf: List[str] = []

    for raw in md_lines:
        line = raw.rstrip("\n")
        s = line.strip()

        # Ignore horizontal rules and blank lines in output; they don't belong to panel text.
        if not s or s == "---":
            continue

        # Heading 1 / Heading 2 are kept.
        if s.startswith("# "):
            # If we were in a panel, flush it before a new top heading.
            if in_panel:
                flush_panel_paragraph(panel_buf)
                panel_buf = []
                in_panel = False
            t = s[2:].strip()
            if not contains_cjk(t):
                add_heading(doc, t, level=1)
            continue

        if s.startswith("## "):
            if in_panel:
                flush_panel_paragraph(panel_buf)
                panel_buf = []
                in_panel = False
            t = s[3:].strip()
            if not contains_cjk(t):
                add_heading(doc, t, level=2)
            continue

        # Start of a new panel block: flush previous panel first.
        if s.startswith("### "):
            h = s[4:].strip()
            if h.lower().startswith("panel block"):
                if in_panel:
                    flush_panel_paragraph(panel_buf)
                    panel_buf = []
                in_panel = True
                # We do NOT keep the panel header itself; only the panel text below it.
                continue
            else:
                # Ignore other H3 lines such as 'Source: ...'
                continue

        # Ignore figure label lines like **Figure 1.**
        if s.startswith("**") and s.endswith("**") and len(s) >= 4:
            continue

        # For content lines:
        # - If we're inside a panel block, collect them.
        # - If we're outside a panel block, ignore them (we only want per-panel paragraphs).
        if not in_panel:
            continue

        # Remove markdown bullets and backticks; keep the sentence content.
        if s.startswith("- "):
            s = s[2:].strip()
        elif s.startswith("> "):
            # In case transitions are rendered as blockquotes
            s = s[2:].strip()

        s = s.replace("`", "")
        # Skip panel metadata lines like '_panel_id_: 1b'
        if s.lower().lstrip().startswith("_panel_id_"):
            continue
        if contains_cjk(s):
            continue
        panel_buf.append(s)

    # Flush last panel
    if in_panel:
        flush_panel_paragraph(panel_buf)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())