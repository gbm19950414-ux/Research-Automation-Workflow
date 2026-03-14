#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Render a FIGURES-only DOCX:
- Insert figure images (png/jpg/tif...) into DOCX
- Add figure titles + figure legends

Design goals:
1) Minimal changes vs your 007_render_docx.py style (template-driven, token replacement ready)
2) Future format switching: image extensions & naming & paths can change without touching code
3) Preferred input: render figures from 08_manuscript/IR/results.ir.yaml; if it has no figure blocks, fall back to 08_manuscript/templates/figures.config.yaml for a minimal figure list + image patterns.

Default project layout (auto-detected by walking up from cwd):
  <ROOT>/
    06_figures/figs/figure_1*.png     (exported figures)
    08_manuscript/IR/*.ir.yaml
    08_manuscript/templates/*.yaml
    08_manuscript/out/<template>/figures.docx

Usage:
  python 008_render_figures_docx.py
  python 008_render_figures_docx.py --template 08_manuscript/templates/journal_short.yaml
  python 008_render_figures_docx.py --figs-dir 06_figures/figs --out 08_manuscript/out/journal_short/figures.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any

import yaml
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image


# -----------------------
# Token regex (reserved; keep compatible with 007)
# -----------------------
RE_CITE = re.compile(r"\{cite:([^}]+)\}")
RE_XREF = re.compile(r"\{xref:([^}]+)\}")


# -----------------------
# Template (same spirit as 007)
# -----------------------
@dataclass
class Template:
    name: str
    section_map: Dict[str, dict]
    rendering: dict
    citations: dict
    styles: dict
    typography: dict


def find_project_root(start: Path) -> Path:
    """
    Walk up from start to find a directory containing both 06_figures and 08_manuscript.
    """
    start = start.resolve()
    for p in [start] + list(start.parents):
        if (p / "06_figures").is_dir() and (p / "08_manuscript").is_dir():
            return p
    return start


def load_yaml_single_doc(path: Path) -> dict:
    """
    Strict single-doc loader (yaml.safe_load).
    If user accidentally saved multi-doc YAML (---), we raise with a clear hint.
    """
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    txt = path.read_text(encoding="utf-8")
    # detect multi-doc quickly
    if re.search(r"(?m)^---\s*$", txt.strip()):
        raise ValueError(
            f"YAML appears to contain multiple documents (---): {path}\n"
            f"Please convert it to a single YAML document for safe_load, "
            f"or split it into separate files."
        )
    return yaml.safe_load(txt) or {}


def load_template(path: Path) -> Template:
    raw = load_yaml_single_doc(path)
    t = raw.get("template", {}) or {}
    name = t.get("name", "template")
    section_map = (t.get("structure", {}) or {}).get("section_map", {}) or {}
    rendering = t.get("rendering", {}) or {}
    citations = t.get("citations", {}) or {}
    styles = t.get("docx_styles", {}) or {}
    typography = t.get("typography", {}) or {}
    return Template(
        name=name,
        section_map=section_map,
        rendering=rendering,
        citations=citations,
        styles=styles,
        typography=typography,
    )


# -----------------------
# Token helpers (same idea as 007; keep xref/cite ready)
# -----------------------
def parse_cite_keys(cite_payload: str) -> List[str]:
    parts = re.split(r"[;,]", cite_payload)
    return [p.strip() for p in parts if p.strip()]


def format_vancouver_numbers(nums: List[int], compress_ranges: bool = True) -> str:
    if not nums:
        return ""
    nums = sorted(set(nums))
    if not compress_ranges:
        return ",".join(str(n) for n in nums)

    ranges: List[str] = []
    start = prev = nums[0]
    for n in nums[1:]:
        if n == prev + 1:
            prev = n
            continue
        if start == prev:
            ranges.append(str(start))
        elif prev == start + 1:
            ranges.append(f"{start},{prev}")
        else:
            ranges.append(f"{start}-{prev}")
        start = prev = n

    if start == prev:
        ranges.append(str(start))
    elif prev == start + 1:
        ranges.append(f"{start},{prev}")
    else:
        ranges.append(f"{start}-{prev}")

    return ",".join(ranges)


def assign_vancouver_numbers_from_texts(texts: List[str]) -> Dict[str, int]:
    """
    Assign citation numbers by first appearance scanning a list of strings.
    """
    cite_num: Dict[str, int] = {}
    counter = 0
    for text in texts:
        for m in RE_CITE.finditer(text or ""):
            keys = parse_cite_keys(m.group(1))
            for k in keys:
                if k not in cite_num:
                    counter += 1
                    cite_num[k] = counter
    return cite_num


def replace_tokens(
    text: str,
    template: Template,
    cite_num: Dict[str, int],
    fig_map: Dict[str, int],
    tbl_map: Dict[str, int],
) -> str:
    def repl_cite(m: re.Match) -> str:
        keys = parse_cite_keys(m.group(1))
        nums = [cite_num[k] for k in keys if k in cite_num]
        if template.citations.get("sort_numeric", True):
            nums = sorted(nums)
        num_str = format_vancouver_numbers(nums, compress_ranges=template.citations.get("compress_ranges", True))
        left = template.citations.get("bracket_left", "[")
        right = template.citations.get("bracket_right", "]")
        delim = template.citations.get("delimiter", ",")
        if delim != ",":
            num_str = num_str.replace(",", delim)
        return f"{left}{num_str}{right}"

    text = RE_CITE.sub(repl_cite, text or "")

    def repl_xref(m: re.Match) -> str:
        target = (m.group(1) or "").strip()
        if target in fig_map:
            pref = template.rendering.get("figure_caption_prefix", "Fig.")
            return f"{pref} {fig_map[target]}"
        if target in tbl_map:
            pref = template.rendering.get("table_caption_prefix", "Table")
            return f"{pref} {tbl_map[target]}"
        return target

    text = RE_XREF.sub(repl_xref, text)
    return text

# -----------------------
# Supplement panel label normalization (render-time only)
# -----------------------
def normalize_supplement_panel_labels_in_text(text: str) -> str:
    """Normalize supplement-style panel labels in rendered legend text.

    Examples:
      (s1_a) -> (a)
      (S2_b) -> (b)
      (s3c)  -> (c)

    This is intentionally applied only at render time for supplement figure output,
    so the underlying IR/YAML structure does not need to change.
    """
    if not isinstance(text, str) or not text:
        return text

    # Parenthesized labels: (s1_a) / (s2b) -> (a) / (b)
    text = re.sub(r"\(([sS]\d+_?([a-zA-Z]))\)", lambda m: f"({m.group(2).lower()})", text)

    # Also normalize comma-separated grouped labels if they ever appear, e.g. (s1_a,s1_b)
    def _normalize_group(match: re.Match) -> str:
        inner = match.group(1)
        parts = [p.strip() for p in inner.split(",")]
        out = []
        for p in parts:
            m = re.fullmatch(r"[sS]\d+_?([a-zA-Z])", p)
            out.append(m.group(1).lower() if m else p)
        return "(" + ",".join(out) + ")"

    text = re.sub(r"\(((?:[sS]\d+_?[a-zA-Z]\s*,\s*)+[sS]\d+_?[a-zA-Z])\)", _normalize_group, text)
    return text


def add_paragraph_with_style(doc: Document, text: str, style_name: str, font_pt: Optional[int] = None):
    p = doc.add_paragraph(text or "")
    if style_name:
        try:
            p.style = doc.styles[style_name]
        except Exception:
            pass
    if font_pt is not None:
        for run in p.runs:
            run.font.size = Pt(font_pt)
    return p


# -----------------------
# DOCX page/image sizing helpers
# -----------------------
def emu_to_inches(emu: int) -> float:
    return float(emu) / 914400.0


def get_writable_page_box_in(doc: Document) -> Tuple[float, float]:
    """Return writable page width/height in inches for the first section."""
    sec = doc.sections[0]
    page_w = emu_to_inches(sec.page_width)
    page_h = emu_to_inches(sec.page_height)
    left = emu_to_inches(sec.left_margin)
    right = emu_to_inches(sec.right_margin)
    top = emu_to_inches(sec.top_margin)
    bottom = emu_to_inches(sec.bottom_margin)
    return max(page_w - left - right, 1.0), max(page_h - top - bottom, 1.0)


def probe_image_size_in(img_path: Path) -> Tuple[float, float]:
    """Read image size and return width/height in inches.

    Uses embedded DPI if available; otherwise falls back to 72 DPI.
    Aspect ratio remains correct either way.
    """
    with Image.open(str(img_path)) as im:
        px_w, px_h = im.size
        dpi = im.info.get("dpi", (72, 72))
        if isinstance(dpi, tuple) and len(dpi) >= 2:
            dpi_x = float(dpi[0] or 72)
            dpi_y = float(dpi[1] or 72)
        else:
            dpi_x = dpi_y = 72.0
        if dpi_x <= 0:
            dpi_x = 72.0
        if dpi_y <= 0:
            dpi_y = 72.0
        return px_w / dpi_x, px_h / dpi_y


def fit_image_size_in(
    img_path: Path,
    preferred_width_in: float,
    max_width_in: float,
    max_height_in: float,
) -> Tuple[float, float]:
    """Return width/height in inches scaled to fit both width and height limits."""
    nat_w, nat_h = probe_image_size_in(img_path)
    if nat_w <= 0 or nat_h <= 0:
        w = min(preferred_width_in, max_width_in)
        return w, w

    # Start from preferred width but never exceed writable width.
    scale = min(preferred_width_in / nat_w, max_width_in / nat_w, 1.0e9)
    w = nat_w * scale
    h = nat_h * scale

    # If still too tall, scale down by height.
    if h > max_height_in:
        h_scale = max_height_in / nat_h
        w = nat_w * h_scale
        h = nat_h * h_scale

    # Final width clamp safety.
    if w > max_width_in:
        w_scale = max_width_in / nat_w
        w = nat_w * w_scale
        h = nat_h * w_scale

    return max(w, 0.5), max(h, 0.5)


# -----------------------
# Figure IR builder (strict)
# -----------------------
def extract_fig_blocks_from_ir(ir: dict) -> List[dict]:
    """
    Return a list of figure blocks from IR (any section).
    Expected block schema example:
      - type: figure
        id: "Figure 1" or "fig1"
        title: "..."
        caption: "..."   # (optional)
        legend: "..."    # (optional)
        image_path: "06_figures/figs/figure_1.png"  # (optional)
        images: ["..."]  # (optional)
    """
    doc = ir.get("document", {}) or {}
    sections = doc.get("sections", []) or []
    fig_blocks: List[dict] = []
    for sec in sections:
        for b in (sec.get("blocks", []) or []):
            if (b.get("type") or "").lower() == "figure":
                fig_blocks.append(b)
    return fig_blocks




def normalize_figure_number(fig_id: str) -> Optional[int]:
    """
    Parse "Figure 1" / "Fig. 1" / "1" / "figure_1" -> 1
    """
    if not fig_id:
        return None
    m = re.search(r"(\d+)", str(fig_id))
    return int(m.group(1)) if m else None


# New helper for IR normalization
from typing import List
def normalize_ir_figures(figs: List[dict]) -> List[dict]:
    """Normalize figures_ir.yaml / supplement_figures_ir.yaml entries to renderer blocks."""
    figure_blocks: List[dict] = []
    for f in figs:
        assets = f.get("assets") or {}
        img = (
            f.get("render_png")
            or assets.get("render_png")
            or assets.get("main_figure_file")
            or f.get("main_figure_file")
        )
        images = [img] if img else []
        figure_blocks.append({
            "type": "figure",
            "id": f.get("figure_id") or f.get("id"),
            "title": f.get("figure_title_en") or "",
            "legend": f.get("figure_legend_en") or "",
            "images": images,
        })
    return figure_blocks


# -----------------------
# Figures config (minimal fallback)
# -----------------------

def load_figures_config(path: Path) -> dict:
    cfg = load_yaml_single_doc(path)
    if not isinstance(cfg, dict):
        raise ValueError(f"Invalid figures config (expected mapping): {path}")
    return cfg


def apply_figures_config_overrides(template: Template, cfg: dict) -> None:
    """Merge rendering-level overrides from figures.config.yaml into template.rendering."""
    r = cfg.get("rendering", {}) or {}
    if not isinstance(r, dict):
        return

    # map config keys -> template.rendering keys
    mapping = {
        "page_break_between_figures": "page_break_between_figures",
        "image_width_in": "figure_image_width_in",
        "image_prefer_ext": "figure_image_prefer_ext",
        "figure_caption_prefix": "figure_caption_prefix",
    }
    for k_cfg, k_tpl in mapping.items():
        if k_cfg in r and r[k_cfg] is not None:
            template.rendering[k_tpl] = r[k_cfg]


def build_figure_blocks_from_config(cfg: dict) -> List[dict]:
    """Create minimal figure blocks from figures.config.yaml.

    Each enabled entry becomes:
      {type: figure, id: 'Figure N', title: '', legend: '', images: [pattern]}

    Note: `images` is used to pass a glob pattern; renderer will resolve it inside figs_dir.
    """
    items = cfg.get("figures") or []
    if not isinstance(items, list):
        raise ValueError("figures.config.yaml: `figures` must be a list")

    blocks: List[dict] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        if it.get("enable") is False:
            continue
        fig_id = str(it.get("id") or "").strip()
        if not fig_id:
            continue

        img = it.get("image", {}) or {}
        pattern = ""
        if isinstance(img, dict):
            pattern = str(img.get("pattern") or "").strip()

        b = {
            "type": "figure",
            "id": fig_id,
            "title": str(it.get("title") or "").strip(),
            "legend": str(it.get("legend") or "").strip(),
        }
        if pattern:
            # store as a pattern hint; renderer will interpret it relative to figs_dir
            b["images"] = [pattern]
        blocks.append(b)

    return blocks


# -----------------------
# Image locating
# -----------------------
def find_figure_image(figs_dir: Path, fig_n: int, prefer_ext: List[str]) -> Optional[Path]:
    """
    Search for figure image files under figs_dir using common patterns.
    Works for your naming like:
      figure_1_画板 1.png
      figure_1.png
      figure_1-anything.jpg
    """
    patterns = [
        f"figure_{fig_n}*.png",
        f"figure_{fig_n}*.jpg",
        f"figure_{fig_n}*.jpeg",
        f"figure_{fig_n}*.tif",
        f"figure_{fig_n}*.tiff",
        f"Figure_{fig_n}*.png",
        f"Figure_{fig_n}*.jpg",
        f"Figure_{fig_n}*.jpeg",
        f"Figure {fig_n}*.png",
        f"Figure {fig_n}*.jpg",
        f"Figure {fig_n}*.jpeg",
    ]

    candidates: List[Path] = []
    for pat in patterns:
        candidates.extend(sorted(figs_dir.glob(pat)))

    # de-dup
    seen = set()
    uniq = []
    for c in candidates:
        if c.resolve() in seen:
            continue
        seen.add(c.resolve())
        uniq.append(c)
    candidates = uniq

    if not candidates:
        return None

    # rank by preferred extension order
    def rank(p: Path) -> Tuple[int, int]:
        ext = p.suffix.lower().lstrip(".")
        try:
            e_rank = prefer_ext.index(ext)
        except ValueError:
            e_rank = 999
        # shorter name first (usually less messy)
        return (e_rank, len(p.name))

    candidates.sort(key=rank)
    return candidates[0]


# -----------------------
# Render
# -----------------------
def render_figures_docx(
    figure_blocks: List[dict],
    template: Template,
    out_path: Path,
    figs_dir: Path,
    figure_mode: str = "main",
):
    doc = Document()

    normal_font_pt = int(template.typography.get("normal_font_pt", 11))
    caption_style = template.styles.get("caption", "Caption")
    h1_style = template.styles.get("h1", "Heading 1")
    normal_style = template.styles.get("normal", "Normal")

    if str(figure_mode).strip().lower() == "supplement":
        pref = template.rendering.get("supplement_figure_caption_prefix", "Supplementary Figure")
    else:
        pref = template.rendering.get("figure_caption_prefix", "Figure")
    img_width_in = float(template.rendering.get("figure_image_width_in", 6.5))
    page_break_between_figures = bool(template.rendering.get("page_break_between_figures", True))
    writable_width_in, writable_height_in = get_writable_page_box_in(doc)
    # Reserve some vertical room for the figure heading and optional caption/legend.
    # This keeps the image itself from overflowing a single page.
    image_max_page_fraction = float(template.rendering.get("figure_image_max_page_fraction", 0.78))
    image_max_height_in = template.rendering.get("figure_image_max_height_in", None)
    if image_max_height_in is None:
        max_img_height_in = writable_height_in * image_max_page_fraction
    else:
        max_img_height_in = float(image_max_height_in)
    max_img_width_in = min(img_width_in, writable_width_in)

    # For token replacement readiness (citations/xrefs). For figures-only doc usually empty.
    # Build fig_map by order of blocks.
    fig_map = {}
    for i, b in enumerate(figure_blocks, start=1):
        fid = b.get("id") or f"Figure {i}"
        fig_map[str(fid)] = i
    tbl_map: Dict[str, int] = {}

    # collect texts for citation numbering
    all_texts = []
    for b in figure_blocks:
        all_texts.append(b.get("title", "") or "")
        all_texts.append(b.get("caption", "") or "")
        all_texts.append(b.get("legend", "") or "")
    cite_num = assign_vancouver_numbers_from_texts(all_texts)

    for idx, b in enumerate(figure_blocks, start=1):
        fid = str(b.get("id") or f"Figure {idx}")
        fig_n = normalize_figure_number(fid) or idx

        title = b.get("title", "") or ""
        caption = b.get("caption", "") or ""
        legend = b.get("legend", "") or ""

        # token replacement
        title = replace_tokens(title, template, cite_num, fig_map, tbl_map)
        caption = replace_tokens(caption, template, cite_num, fig_map, tbl_map)
        legend = replace_tokens(legend, template, cite_num, fig_map, tbl_map)
        if str(figure_mode).strip().lower() == "supplement":
            legend = normalize_supplement_panel_labels_in_text(legend)

        # --- Figure heading line ---
        # Prefer: "Figure 1. <title>" (caption style)
        head_line = f"{pref} {fig_n}. {title}".strip()
        add_paragraph_with_style(doc, head_line, caption_style, font_pt=normal_font_pt)

        if caption:
            add_paragraph_with_style(doc, caption.strip(), normal_style, font_pt=normal_font_pt)

        # --- Image ---
        img_path = None
        # 1) explicit paths
        if b.get("image_path"):
            img_path = Path(b["image_path"])
        elif b.get("image"):
            img_path = Path(b["image"])
        elif b.get("images") and isinstance(b.get("images"), list):
            # pick first; may be a concrete path OR a glob pattern
            first = str(b["images"][0])
            if any(ch in first for ch in ["*", "?", "[", "]"]):
                # treat as glob pattern relative to figs_dir
                matches = sorted(figs_dir.glob(first))
                img_path = matches[0] if matches else Path(first)
            else:
                img_path = Path(first)

        if img_path is not None:
            # resolve relative to project root-like usage; if not exists, try figs_dir
            if not img_path.is_absolute():
                cand = (figs_dir / img_path).resolve()
                if cand.exists():
                    img_path = cand
                else:
                    cand2 = (Path.cwd() / img_path).resolve()
                    if cand2.exists():
                        img_path = cand2
        else:
            # 2) auto-locate
            img_path = find_figure_image(
                figs_dir=figs_dir,
                fig_n=fig_n,
                prefer_ext=template.rendering.get("figure_image_prefer_ext", ["png", "jpg", "jpeg", "tif", "tiff"]),
            )

        if img_path is None or not img_path.exists():
            add_paragraph_with_style(
                doc,
                f"[WARN] Missing figure image for Figure {fig_n}. Searched in: {figs_dir}",
                normal_style,
                font_pt=normal_font_pt,
            )
        else:
            # NOTE: python-docx cannot embed .ai/.pdf as an image directly.
            # You already exported PNG -> best practice.
            try:
                fit_w_in, fit_h_in = fit_image_size_in(
                    img_path=img_path,
                    preferred_width_in=img_width_in,
                    max_width_in=max_img_width_in,
                    max_height_in=max_img_height_in,
                )
                doc.add_picture(str(img_path), width=Inches(fit_w_in), height=Inches(fit_h_in))
            except Exception as e:
                add_paragraph_with_style(
                    doc,
                    f"[ERROR] Failed to insert image: {img_path} ({e})",
                    normal_style,
                    font_pt=normal_font_pt,
                )

        # --- Legend (single paragraph, journal-friendly) ---
        if legend.strip():
            add_paragraph_with_style(doc, legend.strip(), normal_style, font_pt=normal_font_pt)

        # page break
        if page_break_between_figures and idx != len(figure_blocks):
            doc.add_page_break()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Rendered FIGURES DOCX: {out_path}")

def run_one_mode(args, figure_mode: str) -> None:
    # root
    root = Path(args.root).resolve() if args.root.strip() else find_project_root(Path.cwd())
    ms_dir = root / "08_manuscript"

    # template
    if args.template.strip():
        template_path = Path(args.template).resolve()
    else:
        cand = ms_dir / "templates" / "journal_short.yaml"
        template_path = cand if cand.exists() else (ms_dir / "templates" / "thesis.yaml")
    template = load_template(template_path)

    # figures config (optional fallback)
    if args.fig_config.strip():
        fig_cfg_path = Path(args.fig_config).resolve()
    else:
        fig_cfg_path = ms_dir / "templates" / "figures.config.yaml"

    fig_cfg = None
    if fig_cfg_path.exists():
        fig_cfg = load_figures_config(fig_cfg_path)
        apply_figures_config_overrides(template, fig_cfg)

    # dirs
    figs_dir = Path(args.figs_dir).resolve() if args.figs_dir.strip() else (root / "06_figures" / "figs")

    # Prefer figure-specific IR first
    fig_ir_path = (ms_dir / "IR" / "figures_ir.yaml").resolve()
    supplement_ir_path = (ms_dir / "IR" / "supplement_figures_ir.yaml").resolve()
    results_ir_path = (ms_dir / "IR" / "results.ir.yaml").resolve()

    if figure_mode == "supplement":
        if supplement_ir_path.exists():
            ir_path = supplement_ir_path
            print(f"[INFO] Using supplement figure IR: {supplement_ir_path}")
        elif results_ir_path.exists():
            ir_path = results_ir_path
            print(f"[INFO] Using results IR (fallback): {results_ir_path}")
        else:
            raise FileNotFoundError(
                f"Neither supplement_figures_ir.yaml nor results.ir.yaml found in {ms_dir / 'IR'}"
            )
    else:
        if fig_ir_path.exists():
            ir_path = fig_ir_path
            print(f"[INFO] Using figure IR: {fig_ir_path}")
        elif results_ir_path.exists():
            ir_path = results_ir_path
            print(f"[INFO] Using results IR (fallback): {results_ir_path}")
        else:
            raise FileNotFoundError(
                f"Neither figures_ir.yaml nor results.ir.yaml found in {ms_dir / 'IR'}"
            )

    # Output
    if args.out.strip():
        base_out = Path(args.out).resolve()
        if str(base_out).lower().endswith(".docx"):
            if figure_mode == "supplement":
                out_path = base_out.with_name(base_out.stem + "_supplement" + base_out.suffix)
            else:
                out_path = base_out
        else:
            default_name = "supplement_figures.docx" if figure_mode == "supplement" else "figures.docx"
            out_path = base_out / default_name
    else:
        default_name = "supplement_figures.docx" if figure_mode == "supplement" else "figures.docx"
        out_path = ms_dir / "out" / template.name / default_name

    ir = load_yaml_single_doc(ir_path)
    # Normalize figure IR schema if using figures_ir.yaml / supplement_figures_ir.yaml
    if ir_path.name in {"figures_ir.yaml", "supplement_figures_ir.yaml"}:
        figs = ir.get("figures") or []
        figure_blocks = normalize_ir_figures(figs)
    else:
        figure_blocks = extract_fig_blocks_from_ir(ir)

    if not figure_blocks:
        if fig_cfg is None:
            raise ValueError(
                f"No figure blocks found in IR: {ir_path}\n"
                f"And figures config not found: {fig_cfg_path}\n"
                f"Provide figure blocks in IR OR create figures.config.yaml."
            )
        figure_blocks = build_figure_blocks_from_config(fig_cfg)
        if not figure_blocks:
            raise ValueError(
                f"IR has no figure blocks and figures.config.yaml produced none: {fig_cfg_path}\n"
                f"Check that `figures:` list entries have id and enable: true."
            )
        print(f"[WARN] IR has no figure blocks; using figures.config.yaml for figure list: {fig_cfg_path}")

    render_figures_docx(
        figure_blocks=figure_blocks,
        template=template,
        out_path=out_path,
        figs_dir=figs_dir,
        figure_mode=figure_mode,
    )
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", default="", help="Project root (contains 06_figures and 08_manuscript). Default: auto-detect from cwd.")
    parser.add_argument("--template", default="", help="Template YAML. Default: 08_manuscript/templates/journal_short.yaml if exists.")
    parser.add_argument("--fig-config", default="", help="Figures config YAML (fallback). Default: 08_manuscript/templates/figures.config.yaml")
    parser.add_argument("--figs-dir", default="", help="06_figures/figs dir (images).")
    parser.add_argument("--out", default="", help="Output DOCX path. Default: 08_manuscript/out/<template_name>/figures.docx")
    parser.add_argument("--figure-mode", choices=["main", "supplement", "all"], default="all", help="Render main figures, supplement figures, or both (default: all)")
    args = parser.parse_args()

    if args.figure_mode == "all":
        modes = ["main", "supplement"]
    else:
        modes = [args.figure_mode]

    for mode in modes:
        run_one_mode(args, mode)

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"[FATAL] {e}", file=sys.stderr)
        sys.exit(1)