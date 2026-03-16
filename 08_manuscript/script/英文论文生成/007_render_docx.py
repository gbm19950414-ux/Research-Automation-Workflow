#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Render a manuscript DOCX from YAML IR + template YAML.

Default paths follow your directory convention:
- IR:        08_manuscript/IR/manuscript.ir.yaml (if missing, auto-select an IR file in 08_manuscript/IR/)
- Template:  08_manuscript/templates/journal_short.yaml
- Output:    08_manuscript/out/<template_name>/manuscript.docx

Supports inline tokens:
- {cite:Key} or {cite:Key1;Key2}  -> Vancouver [n]
- {xref:fig1} / {xref:tbl1}      -> Fig. 1 / Table 1 (prefix from template)

Optional reference database:
- 08_manuscript/yaml/references.yaml
  references:
    Smith2022:
      title: "..."
      authors: "Smith J; ..."
      year: 2022
      journal: "..."
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import sys

import yaml
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# -----------------------
# Paths (fixed to your layout)
# -----------------------

PROJECT_ROOT = Path("/Volumes/Samsung_SSD_990_PRO_2TB_Media/EphB1")
MANUSCRIPT_DIR = PROJECT_ROOT / "08_manuscript"

DEFAULT_IR_PATH = MANUSCRIPT_DIR / "IR" / "manuscript.ir.yaml"
DEFAULT_TEMPLATE_PATH = MANUSCRIPT_DIR / "templates" / "journal_short.yaml"
DEFAULT_REF_DB_PATH = MANUSCRIPT_DIR / "yaml" / "references.yaml"
DEFAULT_OUT_DIR = MANUSCRIPT_DIR / "out"


def resolve_default_ir_path() -> Path:
    if DEFAULT_IR_PATH.exists():
        return DEFAULT_IR_PATH

    ir_dir = MANUSCRIPT_DIR / "IR"
    candidates_yaml = list(ir_dir.glob("*.ir.yaml"))
    candidates_yml = list(ir_dir.glob("*.ir.yml"))
    candidates = candidates_yaml + candidates_yml

    if not candidates:
        raise FileNotFoundError(
            f"Default IR file not found at {DEFAULT_IR_PATH}.\n"
            f"Please place an IR file under {ir_dir}/ or pass --ir argument."
        )

    if len(candidates) == 1:
        return candidates[0]

    # Multiple candidates: prefer manuscript.ir.yaml or manuscript.ir.yml
    preferred_names = {"manuscript.ir.yaml", "manuscript.ir.yml"}
    for pref_name in preferred_names:
        pref_path = ir_dir / pref_name
        if pref_path in candidates:
            return pref_path

    # Otherwise choose the first IR file by filename order (deterministic default)
    candidates.sort(key=lambda p: p.name)
    return candidates[0]

def is_manuscript_ir(path: Path) -> bool:
    return path.name in ("manuscript.ir.yaml", "manuscript.ir.yml")


# -----------------------
# Regex tokens
# -----------------------

RE_CITE = re.compile(r"\{cite:([^}]+)\}")
RE_XREF = re.compile(r"\{xref:([^}]+)\}")


@dataclass
class Template:
    name: str
    section_map: Dict[str, dict]
    rendering: dict
    citations: dict
    styles: dict
    typography: dict


def load_yaml(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_template(path: Path) -> Template:
    raw = load_yaml(path)
    t = raw.get("template", {})
    name = t.get("name", "template")
    section_map = (t.get("structure", {}) or {}).get("section_map", {}) or {}
    rendering = t.get("rendering", {}) or {}
    citations = t.get("citations", {}) or {}
    styles = t.get("docx_styles", {}) or {}
    typography = t.get("typography", {}) or {}
    return Template(
        name=name,
        section_map=section_map,
        rendering=rendering,
        citations=citations,
        styles=styles,
        typography=typography,
    )


def load_reference_db(path: Path) -> Dict[str, dict]:
    if not path.exists():
        return {}
    raw = load_yaml(path)
    return raw.get("references", {}) or {}


def iter_sections(ir: dict) -> List[dict]:
    doc = ir.get("document", {}) or {}
    sections = doc.get("sections", []) or []
    if not isinstance(sections, list):
        raise ValueError("IR document.sections must be a list")
    return sections


def build_figure_table_numbering(sections: List[dict]) -> Tuple[Dict[str, int], Dict[str, int]]:
    fig_map: Dict[str, int] = {}
    tbl_map: Dict[str, int] = {}
    fig_n = 0
    tbl_n = 0

    for sec in sections:
        blocks = sec.get("blocks", []) or []
        for b in blocks:
            btype = b.get("type")
            if btype == "figure":
                fid = b.get("id")
                if fid and fid not in fig_map:
                    fig_n += 1
                    fig_map[fid] = fig_n
            elif btype == "table":
                tid = b.get("id")
                if tid and tid not in tbl_map:
                    tbl_n += 1
                    tbl_map[tid] = tbl_n
    return fig_map, tbl_map


def parse_cite_keys(cite_payload: str) -> List[str]:
    # allow separators ; or , and remove spaces
    parts = re.split(r"[;,]", cite_payload)
    keys = [p.strip() for p in parts if p.strip()]
    return keys


def assign_vancouver_numbers(
    sections: List[dict],
) -> Dict[str, int]:
    """
    Assign citation numbers by first appearance scanning all text-like fields.
    """
    cite_num: Dict[str, int] = {}
    counter = 0

    def scan_text(text: str):
        nonlocal counter
        for m in RE_CITE.finditer(text):
            keys = parse_cite_keys(m.group(1))
            for k in keys:
                if k not in cite_num:
                    counter += 1
                    cite_num[k] = counter

    for sec in sections:
        for b in sec.get("blocks", []) or []:
            btype = b.get("type")
            if btype in ("paragraph", "heading"):
                scan_text(b.get("text", "") or "")
            elif btype == "list":
                for item in (b.get("items", []) or []):
                    scan_text(str(item))
            elif btype in ("figure", "table"):
                scan_text(b.get("caption", "") or "")
                scan_text(b.get("title", "") or "")
    return cite_num


def format_vancouver_numbers(nums: List[int], compress_ranges: bool = True) -> str:
    """
    Given a list of integers, return "1,2,5-7" if compress_ranges else "1,2,5,6,7".
    """
    if not nums:
        return ""
    nums = sorted(set(nums))
    if not compress_ranges:
        return ",".join(str(n) for n in nums)

    ranges: List[str] = []
    start = prev = nums[0]
    for n in nums[1:]:
        if n == prev + 1:
            prev = n
            continue
        # close previous run
        if start == prev:
            ranges.append(str(start))
        elif prev == start + 1:
            ranges.append(f"{start},{prev}")
        else:
            ranges.append(f"{start}-{prev}")
        start = prev = n

    # close last run
    if start == prev:
        ranges.append(str(start))
    elif prev == start + 1:
        ranges.append(f"{start},{prev}")
    else:
        ranges.append(f"{start}-{prev}")

    # ensure commas only
    return ",".join(ranges)


def replace_tokens(
    text: str,
    template: Template,
    cite_num: Dict[str, int],
    fig_map: Dict[str, int],
    tbl_map: Dict[str, int],
) -> str:
    # citations
    def repl_cite(m: re.Match) -> str:
        keys = parse_cite_keys(m.group(1))
        nums = [cite_num[k] for k in keys if k in cite_num]
        if template.citations.get("sort_numeric", True):
            nums = sorted(nums)
        num_str = format_vancouver_numbers(nums, compress_ranges=template.citations.get("compress_ranges", True))
        left = template.citations.get("bracket_left", "[")
        right = template.citations.get("bracket_right", "]")
        delim = template.citations.get("delimiter", ",")
        # format_vancouver_numbers returns comma-separated; honor delimiter if different
        if delim != ",":
            num_str = num_str.replace(",", delim)
        return f"{left}{num_str}{right}"

    text = RE_CITE.sub(repl_cite, text)

    # xrefs
    def repl_xref(m: re.Match) -> str:
        target = (m.group(1) or "").strip()
        # figure?
        if target in fig_map:
            pref = template.rendering.get("figure_caption_prefix", "Fig.")
            return f"{pref} {fig_map[target]}"
        if target in tbl_map:
            pref = template.rendering.get("table_caption_prefix", "Table")
            return f"{pref} {tbl_map[target]}"
        return target  # fallback

    text = RE_XREF.sub(repl_xref, text)
    return text


def add_paragraph_with_style(doc: Document, text: str, style_name: str, font_pt: Optional[int] = None):
    p = doc.add_paragraph(text)
    if style_name:
        try:
            p.style = doc.styles[style_name]
        except Exception:
            # If style missing, leave default
            pass
    if font_pt is not None:
        for run in p.runs:
            run.font.size = Pt(font_pt)
    return p


# --- Front matter helpers ---

def _meta_list(val):
    return val if isinstance(val, list) else []


def _meta_dict(val):
    return val if isinstance(val, dict) else {}


def format_author_name(author: object) -> str:
    if isinstance(author, str):
        return author.strip()
    if not isinstance(author, dict):
        return ""
    name = str(author.get("name") or "").strip()
    if not name:
        return ""
    marks: List[str] = []
    if author.get("equal_contribution"):
        marks.append("†")
    if author.get("corresponding"):
        marks.append("*")
    return name + ("".join(marks) if marks else "")


def format_authors_line(meta: dict) -> str:
    authors = _meta_list(meta.get("authors"))
    parts = [format_author_name(a) for a in authors]
    parts = [p for p in parts if p]
    return ", ".join(parts)


def collect_affiliation_lines(meta: dict) -> List[str]:
    lines: List[str] = []

    affs = _meta_list(meta.get("affiliations"))
    if affs:
        for aff in affs:
            if isinstance(aff, str):
                t = aff.strip()
                if t:
                    lines.append(t)
            elif isinstance(aff, dict):
                aff_id = aff.get("id")
                text = str(aff.get("text") or "").strip()
                if text:
                    if aff_id not in (None, ""):
                        lines.append(f"{aff_id}. {text}")
                    else:
                        lines.append(text)
        return lines

    # Fallback: derive affiliations from author entries if top-level affiliations are absent.
    seen = set()
    for author in _meta_list(meta.get("authors")):
        if not isinstance(author, dict):
            continue
        aff = str(author.get("affiliation") or "").strip()
        if aff and aff not in seen:
            seen.add(aff)
            lines.append(aff)
    return lines


def format_correspondence_line(meta: dict) -> str:
    corr = meta.get("corresponding_author")

    # Case 1: list of corresponding authors
    if isinstance(corr, list):
        parts = []
        for item in corr:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or "").strip()
            email = str(item.get("email") or "").strip()
            if name and email:
                parts.append(f"{name} ({email})")
            elif email:
                parts.append(email)
            elif name:
                parts.append(name)
        if parts:
            return "Correspondence: " + "; ".join(parts)
        return ""

    # Case 2: single corresponding author dict
    if isinstance(corr, dict):
        name = str(corr.get("name") or "").strip()
        email = str(corr.get("email") or "").strip()
        if name and email:
            return f"Correspondence: {name} ({email})"
        if email:
            return f"Correspondence: {email}"
        if name:
            return f"Correspondence: {name}"

    return ""

def collect_footnote_lines(meta: dict) -> List[str]:
    foot = _meta_dict(meta.get("footnotes"))
    out: List[str] = []
    for _, v in foot.items():
        txt = str(v or "").strip()
        if txt:
            out.append(txt)
    return out


def render_front_matter(doc: Document, template: Template, meta: dict):
    """Render title-page metadata in a plain, submission-style layout.

    Layout:
    - Title (centered, existing title style)
    - Author line (centered, normal style)
    - Affiliation lines (centered, normal style)
    - Correspondence line (centered, normal style)
    - Footnotes (left-aligned, normal style)
    """
    normal_style = template.styles.get("normal", "Normal")
    normal_font_pt = int(template.typography.get("normal_font_pt", 11))

    title = str(meta.get("title") or "").strip()
    if title:
        p = add_paragraph_with_style(doc, title, template.styles.get("title_style", "Title"), font_pt=None)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    authors_line = format_authors_line(meta)
    if authors_line:
        p = add_paragraph_with_style(doc, authors_line, normal_style, font_pt=normal_font_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for aff_line in collect_affiliation_lines(meta):
        p = add_paragraph_with_style(doc, aff_line, normal_style, font_pt=normal_font_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    corr_line = format_correspondence_line(meta)
    if corr_line:
        p = add_paragraph_with_style(doc, corr_line, normal_style, font_pt=normal_font_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for foot_line in collect_footnote_lines(meta):
        add_paragraph_with_style(doc, foot_line, normal_style, font_pt=normal_font_pt)

    if title or authors_line or collect_affiliation_lines(meta) or corr_line or collect_footnote_lines(meta):
        doc.add_paragraph("")


def render(ir_path: Path, template_path: Path, out_path: Path, ref_db_path: Path):
    print(f"[INFO] Using IR: {ir_path}")
    ir = load_yaml(ir_path)
    template = load_template(template_path)
    ref_db = load_reference_db(ref_db_path)

    sections = iter_sections(ir)

    fig_map, tbl_map = build_figure_table_numbering(sections)
    cite_num = assign_vancouver_numbers(sections)

    doc = Document()

    normal_font_pt = int(template.typography.get("normal_font_pt", 11))

    # Manuscript front matter from document.meta
    meta = (ir.get("document", {}) or {}).get("meta", {}) or {}
    render_front_matter(doc, template, meta)

    # render sections
    for sec in sections:
        sec_id = sec.get("id")
        sec_title = sec.get("title", sec_id)

        map_entry = template.section_map.get(sec_id, {"title": sec_title, "include": True})
        if not map_entry.get("include", True):
            continue

        display_title = map_entry.get("title", sec_title)

        # Section heading (H1)
        # Skip a visible heading for the title section because front matter is rendered above.
        if sec_id != "title":
            add_paragraph_with_style(doc, display_title, template.styles.get("h1", "Heading 1"), font_pt=None)

        for b in sec.get("blocks", []) or []:
            btype = b.get("type")

            if btype == "heading":
                lvl = int(b.get("level", 2))
                style = template.styles.get({2: "h2", 3: "h3"}.get(lvl, "h2"), "Heading 2")
                text = replace_tokens(b.get("text", "") or "", template, cite_num, fig_map, tbl_map)
                add_paragraph_with_style(doc, text, style, font_pt=None)

            elif btype == "paragraph":
                text = replace_tokens(b.get("text", "") or "", template, cite_num, fig_map, tbl_map)
                add_paragraph_with_style(doc, text, template.styles.get("normal", "Normal"), font_pt=normal_font_pt)

            elif btype == "list":
                ordered = bool(b.get("ordered", False))
                items = b.get("items", []) or []
                for i, item in enumerate(items, start=1):
                    t = replace_tokens(str(item), template, cite_num, fig_map, tbl_map)
                    prefix = f"{i}. " if ordered else "• "
                    add_paragraph_with_style(doc, prefix + t, template.styles.get("normal", "Normal"), font_pt=normal_font_pt)

            elif btype == "figure":
                fid = b.get("id")
                caption = b.get("caption", "") or ""
                caption = replace_tokens(caption, template, cite_num, fig_map, tbl_map)
                pref = template.rendering.get("figure_caption_prefix", "Fig.")
                num = fig_map.get(fid, "?")
                cap_line = f"{pref} {num}. {caption}".strip()
                add_paragraph_with_style(doc, cap_line, template.styles.get("caption", "Caption"), font_pt=normal_font_pt)

            elif btype == "table":
                tid = b.get("id")
                caption = b.get("caption", "") or ""
                caption = replace_tokens(caption, template, cite_num, fig_map, tbl_map)
                pref = template.rendering.get("table_caption_prefix", "Table")
                num = tbl_map.get(tid, "?")
                cap_line = f"{pref} {num}. {caption}".strip()
                add_paragraph_with_style(doc, cap_line, template.styles.get("caption", "Caption"), font_pt=normal_font_pt)

            else:
                # unknown block type: ignore safely
                continue

        if template.rendering.get("page_break_between_sections", False):
            doc.add_page_break()

    # References
    if template.section_map.get("references", {"include": True}).get("include", True) and cite_num:
        heading = template.rendering.get("references_heading", "References")
        add_paragraph_with_style(doc, heading, template.styles.get("h1", "Heading 1"), font_pt=None)

        # Sort by assigned number
        items = sorted(cite_num.items(), key=lambda kv: kv[1])
        for key, n in items:
            ref = ref_db.get(key, {})
            # Simple Vancouver-like fallback line
            if ref:
                # Prefer a compact single-line representation
                authors = ref.get("authors", "")
                year = ref.get("year", "")
                title_r = ref.get("title", "") or key
                journal = ref.get("journal", "")
                line_parts = [p for p in [authors, str(year) if year else "", title_r, journal] if p]
                ref_text = ". ".join(line_parts).strip()
            else:
                ref_text = key

            add_paragraph_with_style(doc, f"[{n}] {ref_text}", template.styles.get("reference", "Normal"), font_pt=normal_font_pt)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Rendered DOCX: {out_path}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--ir", default=None, help="Path to manuscript IR YAML")
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE_PATH), help="Path to template YAML")
    parser.add_argument("--out", default="", help="Output docx path (optional)")
    parser.add_argument("--refdb", default=str(DEFAULT_REF_DB_PATH), help="Path to references.yaml (optional)")
    args = parser.parse_args()

    if args.ir and args.ir.strip():
        ir_path = Path(args.ir)
    else:
        try:
            ir_path = resolve_default_ir_path()
        except FileNotFoundError as e:
            print(f"[ERROR] {e}", file=sys.stderr)
            sys.exit(1)

    # Implement default mode switch for assembled manuscript IR
    if args.template == str(DEFAULT_TEMPLATE_PATH) and is_manuscript_ir(ir_path):
        template_path = MANUSCRIPT_DIR / "templates" / "thesis.yaml"
        print("[INFO] Assembled manuscript IR detected; defaulting template to thesis.yaml")
    else:
        template_path = Path(args.template)

    ref_db_path = Path(args.refdb)

    t = load_template(template_path)
    if args.out.strip():
        out_path = Path(args.out)
    else:
        out_path = DEFAULT_OUT_DIR / t.name / "manuscript.docx"

    render(ir_path, template_path, out_path, ref_db_path)


if __name__ == "__main__":
    main()