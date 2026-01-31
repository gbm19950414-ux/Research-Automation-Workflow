#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Export methods from method.yaml to a .docx file.

Behavior:
- For each item under top-level key `methods`:
  - Use `title` as Heading 1
  - Concatenate all `sentences[*].text` (in order) into ONE paragraph
  - Exclude sentences where required is explicitly false
- Skip items with no sentences (or empty sentences)
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any, Dict, List

import yaml
from docx import Document


def normalize_text(s: str) -> str:
    """Normalize YAML folded text into clean single-space sentences."""
    if s is None:
        return ""
    # Replace line breaks/tabs with spaces
    s = re.sub(r"[\r\n\t]+", " ", s)
    # Collapse multiple spaces
    s = re.sub(r"\s{2,}", " ", s)
    return s.strip()


def load_yaml(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def export_to_docx(methods: List[Dict[str, Any]], out_path: Path) -> None:
    doc = Document()

    for m in methods:
        title = (m.get("title") or m.get("id") or "Untitled").strip()
        sentences = m.get("sentences")

        # Some method entries may not have sentences (e.g., placeholder sections)
        if not sentences or not isinstance(sentences, list):
            continue

        # Collect sentence texts in order
        parts: List[str] = []
        for sent in sentences:
            if not isinstance(sent, dict):
                continue

            # Exclude optional sentences explicitly marked as required: false
            if sent.get("required") is False:
                continue

            text = normalize_text(str(sent.get("text") or ""))
            if text:
                parts.append(text)

        if not parts:
            continue

        paragraph_text = " ".join(parts).strip()

        # Title as Heading 1
        doc.add_heading(title, level=1)
        # One paragraph per title
        doc.add_paragraph(paragraph_text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def build_methods_ir(methods: List[Dict[str, Any]]) -> Dict[str, Any]:
    sections_blocks = []
    for m in methods:
        title = (m.get("title") or m.get("id") or "Untitled").strip()
        sentences = m.get("sentences")

        if not sentences or not isinstance(sentences, list):
            continue

        parts: List[str] = []
        for sent in sentences:
            if not isinstance(sent, dict):
                continue
            if sent.get("required") is False:
                continue
            text = normalize_text(str(sent.get("text") or ""))
            if text:
                parts.append(text)

        if not parts:
            continue

        paragraph_text = " ".join(parts).strip()

        sections_blocks.append(
            {"type": "heading", "level": 2, "text": title}
        )
        sections_blocks.append(
            {"type": "paragraph", "text": paragraph_text}
        )

    ir = {
        "ir_version": "0.1",
        "document": {
            "meta": {
                "id": "ephb1_methods",
                "language": "en",
                "title": "",
                "authors": [],
                "date": "",
            },
            "sections": [
                {
                    "id": "methods",
                    "title": "Methods",
                    "blocks": sections_blocks,
                }
            ],
        },
    }
    return ir


def write_ir_yaml(ir: Dict[str, Any], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as f:
        yaml.safe_dump(ir, f, sort_keys=False, allow_unicode=True, width=1000)


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Export method.yaml to a DOCX file (one title = one paragraph) and Methods IR YAML."
    )
    ap.add_argument(
        "-i",
        "--input",
        default="08_manuscript/yaml/method.yaml",
        help="Path to method.yaml (default: 08_manuscript/yaml/method.yaml)",
    )
    ap.add_argument(
        "-o",
        "--output",
        default="08_manuscript/out/methods_export.docx",
        help="Output .docx path (default: 08_manuscript/out/methods_export.docx)",
    )
    ap.add_argument(
        "--out_ir",
        default="08_manuscript/IR/methods.ir.yaml",
        help="Output Methods IR YAML path (default: 08_manuscript/IR/methods.ir.yaml)",
    )
    args = ap.parse_args()

    in_path = Path(args.input).expanduser().resolve()
    out_path = Path(args.output).expanduser().resolve()
    out_ir_path = Path(args.out_ir).expanduser().resolve()

    if not in_path.exists():
        raise FileNotFoundError(f"Input YAML not found: {in_path}")

    data = load_yaml(in_path)
    methods = data.get("methods")
    if not methods or not isinstance(methods, list):
        raise ValueError("YAML must contain a top-level `methods:` list.")

    export_to_docx(methods, out_path)
    ir = build_methods_ir(methods)
    write_ir_yaml(ir, out_ir_path)

    print(f"Saved DOCX: {out_path}")
    print(f"Saved IR: {out_ir_path}")


if __name__ == "__main__":
    main()