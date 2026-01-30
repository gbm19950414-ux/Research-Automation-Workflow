#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Export English abstract from 08_manuscript/abstract.yaml
to 08_manuscript/abstract_en.docx
and also generate IR YAML file for later pipeline use.

No arguments required.
"""

from pathlib import Path
import yaml
from docx import Document
from docx.shared import Pt
import re


# =========================
# 固定路径配置
# =========================

PROJECT_ROOT = Path("/Volumes/Samsung_SSD_990_PRO_2TB_Media/EphB1")
MANUSCRIPT_DIR = PROJECT_ROOT / "08_manuscript"

YAML_PATH = MANUSCRIPT_DIR / "yaml" / "abstract_en.yaml"
OUT_PATH = MANUSCRIPT_DIR / "docx" / "abstract_en.docx"
OUT_IR_PATH = MANUSCRIPT_DIR / "IR" / "abstract.ir.yaml"

FIELDS_ORDER = [
    "background_gap",
    "research_question",
    "approach",
    "key_finding",
    "implication",
]


# =========================
# 核心逻辑
# =========================

def load_yaml(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"abstract.yaml not found: {path}")
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def split_bilingual_block(text: str) -> tuple[str, str]:
    # Split into non-empty lines
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    en_lines = []
    cn_lines = []
    cjk_re = re.compile(r"[\u4e00-\u9fff]")
    chinese_started = False
    for line in lines:
        if not chinese_started:
            if cjk_re.search(line):
                chinese_started = True
                cn_lines.append(line)
            else:
                en_lines.append(line)
        else:
            cn_lines.append(line)
    en = " ".join(en_lines).strip()
    cn = " ".join(cn_lines).strip()
    return en, cn


def build_abstract_en(data: dict) -> str:
    abstract = data.get("abstract", {})
    parts = []

    for key in FIELDS_ORDER:
        node = abstract.get(key)
        if isinstance(node, dict):
            if "en" not in node:
                raise ValueError(f"Missing English abstract field: abstract.{key}.en")
            text = node["en"].strip()
            if not text:
                raise ValueError(f"Empty content in abstract.{key}.en")
        elif isinstance(node, str):
            en, _ = split_bilingual_block(node)
            if not en:
                raise ValueError(f"Empty English content in abstract.{key}")
            text = en
        else:
            raise ValueError(f"Invalid abstract field type for abstract.{key}")
        parts.append(text)

    # 期刊版：单段落，用空格拼接
    return " ".join(parts)


def build_abstract_ir(data: dict) -> dict:
    abstract_en = build_abstract_en(data)
    ir = {
        "ir_version": "0.1",
        "document": {
            "meta": {
                "id": "ephb1_abstract",
                "language": "en",
                "title": "",
                "authors": [],
                "date": "",
            },
            "sections": [
                {
                    "id": "abstract",
                    "title": "Abstract",
                    "blocks": [
                        {
                            "type": "paragraph",
                            "text": abstract_en,
                        }
                    ],
                }
            ],
        },
    }
    return ir


def write_docx(text: str, out_path: Path) -> None:
    doc = Document()

    # Title
    title = doc.add_paragraph("Abstract")
    title.style = doc.styles["Title"]

    # Body
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)

    doc.save(out_path)


def write_ir_yaml(ir: dict, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as f:
        yaml.safe_dump(ir, f, sort_keys=False, allow_unicode=True, width=1000)


# =========================
# 入口
# =========================

def main():
    print("[INFO] Loading abstract.yaml")
    data = load_yaml(YAML_PATH)

    print("[INFO] Building English abstract")
    abstract_en = build_abstract_en(data)

    print("[INFO] Writing abstract_en.docx")
    write_docx(abstract_en, OUT_PATH)

    print("[INFO] Building IR YAML")
    ir = build_abstract_ir(data)

    print(f"[INFO] Writing IR YAML to {OUT_IR_PATH}")
    write_ir_yaml(ir, OUT_IR_PATH)

    print(f"[OK] Exported DOCX: {OUT_PATH}")
    print(f"[OK] Exported IR YAML: {OUT_IR_PATH}")


if __name__ == "__main__":
    main()