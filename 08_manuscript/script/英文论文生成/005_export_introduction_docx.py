#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
export_introduction_docx.py

Default behavior (no args):
- Read: 08_manuscript/introduction_en.yaml (same folder as this script)
- Write: 08_manuscript/Introduction.docx
- Output only submission-ready Introduction text:
  - Keep only: "Introduction" heading + paragraphs in order
  - Paragraph = topic_sentence + sentences[] merged
  - Do NOT include YAML scaffolding fields: source_trace, pid, bridge, hypotheses, narrative_role, etc.
  - Do NOT include subsection headings by default (Nature-style)
  - Citations in `citations: ["[1]", ...]` are appended to paragraph end if not already present
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def _load_yaml(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    if not isinstance(data, dict):
        raise ValueError("YAML root must be a mapping (dict).")
    return data


def _normalize_space(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+\n", "\n", text)
    return text.strip()


def _format_paragraph(topic_sentence: Optional[str], sentences: Optional[List[str]]) -> str:
    parts: List[str] = []
    if isinstance(topic_sentence, str) and topic_sentence.strip():
        parts.append(topic_sentence.strip())
    if isinstance(sentences, list):
        parts.extend([s.strip() for s in sentences if isinstance(s, str) and s.strip()])
    return _normalize_space(" ".join(parts))


def _append_citations(par_text: str, citations: Optional[List[str]]) -> str:
    """Append citations at paragraph end.

    Supports two input formats in YAML:
      1) Bracket indices: [1], [7-9], [1,2]  -> appended verbatim (legacy)
      2) DOI strings: 10.xxxx/..., doi:10.xxxx/..., https://doi.org/10.xxxx/... -> appended as {doi:...}

    If a citation token already appears in the paragraph text, it will not be duplicated.
    """
    if not citations or not isinstance(citations, list):
        return par_text

    doi_pat = re.compile(r"10\.\d{4,9}/[^\s<>\"'\]\)；;，,]+", re.IGNORECASE)

    def _to_doi_token(raw: str) -> Optional[str]:
        s = (raw or "").strip()
        if not s:
            return None

        # Legacy bracket citations
        if re.fullmatch(r"\[[0-9,\-\s]+\]", s):
            return s

        # DOI in various forms
        m = doi_pat.search(s)
        if not m:
            return None
        doi = m.group(0).rstrip(".")
        return "{doi:" + doi + "}"

    cleaned: List[str] = []
    for c in citations:
        if isinstance(c, str):
            tok = _to_doi_token(c)
            if tok:
                cleaned.append(tok)

    if not cleaned:
        return par_text

    missing = [c for c in cleaned if c not in par_text]
    if not missing:
        return par_text

    # Nature-style: citations follow the sentence-ending punctuation.
    if par_text.endswith((".", "!", "?")):
        return f"{par_text} {' '.join(missing)}"
    else:
        return f"{par_text}. {' '.join(missing)}"


def _iter_blocks(data: Dict[str, Any]) -> List[Tuple[Optional[str], Dict[str, Any]]]:
    # Accept either top-level or under manuscript:
    manuscript = data.get("manuscript", data)
    sections = manuscript.get("sections")
    if not isinstance(sections, list):
        raise ValueError("Expected manuscript.sections to be a list.")

    blocks: List[Tuple[Optional[str], Dict[str, Any]]] = []
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        heading = sec.get("heading") if isinstance(sec.get("heading"), str) else None
        paragraphs = sec.get("paragraphs", [])
        if not isinstance(paragraphs, list):
            continue
        for p in paragraphs:
            if isinstance(p, dict):
                blocks.append((heading, p))
    return blocks


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    # East Asia font fallback
    rpr = style.element.rPr
    rFonts = rpr.rFonts
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_docx(
    yaml_path: Path,
    out_path: Path,
    include_subheadings: bool = False,  # keep False for Nature-like formatting
    keep_hypotheses: bool = False,      # drop hypotheses by default for submission
) -> None:
    data = _load_yaml(yaml_path)

    doc = Document()
    _set_doc_defaults(doc)

    # Title
    doc.add_heading("Introduction", level=1)

    blocks = _iter_blocks(data)
    last_heading: Optional[str] = None

    for heading, p in blocks:
        # Drop hypothesis blocks by default (your YAML stores them under `hypotheses`)
        if (not keep_hypotheses) and ("hypotheses" in p):
            continue

        # Do not include subheadings unless explicitly enabled
        if include_subheadings and heading and heading != last_heading:
            doc.add_heading(heading, level=2)
            last_heading = heading

        topic_sentence = p.get("topic_sentence")
        sentences = p.get("sentences")
        par_text = _format_paragraph(topic_sentence, sentences)

        # Append citations from YAML field
        citations = p.get("citations")
        par_text = _append_citations(par_text, citations)

        if par_text:
            doc.add_paragraph(par_text)

    doc.save(str(out_path))


def main() -> None:
    here = Path(__file__).resolve().parent
    yaml_path = here / "introduction_en.yaml"
    out_path = here / "Introduction.docx"

    if not yaml_path.exists():
        raise FileNotFoundError(f"YAML not found: {yaml_path}")

    build_docx(
        yaml_path=yaml_path,
        out_path=out_path,
        include_subheadings=False,
        keep_hypotheses=False,
    )
    print(f"OK: wrote {out_path}")


if __name__ == "__main__":
    main()