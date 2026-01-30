#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
md_to_docx.py
Convert a Markdown file to a .docx file (Word), using python-docx.

Features (pragmatic, not a full Markdown spec):
- Headings: # .. ######  -> Word Heading 1..6
- Paragraphs
- Bullet lists: "-", "*", "+" (including simple indentation)
- Numbered lists: "1. " "2. "
- Blockquotes: lines starting with ">"
- Fenced code blocks: ```lang ... ``` (monospace style)
- Simple tables: pipe tables (best-effort)
- Optionally strip HTML comment blocks, including <!-- TRACE ... --> blocks

Usage:
  python md_to_docx.py input.md -o output.docx
  python md_to_docx.py /mnt/data/Discussion_T8_final_en.md -o Discussion_for_review.docx --strip-trace

Requirements:
  pip install python-docx
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn


TRACE_BLOCK_RE = re.compile(r"<!--\s*TRACE\b.*?-->\s*", re.DOTALL | re.IGNORECASE)
HTML_COMMENT_RE = re.compile(r"<!--.*?-->\s*", re.DOTALL)


@dataclass
class Table:
    header: List[str]
    rows: List[List[str]]


def strip_trace_blocks(text: str, strip_all_comments: bool = False) -> str:
    """
    Remove TRACE blocks (and optionally all HTML comments).
    """
    text = TRACE_BLOCK_RE.sub("", text)
    if strip_all_comments:
        text = HTML_COMMENT_RE.sub("", text)
    return text


def is_table_separator(line: str) -> bool:
    # e.g. |---|:---:|---|
    s = line.strip()
    if "|" not in s:
        return False
    s = s.strip("|").strip()
    if not s:
        return False
    cells = [c.strip() for c in s.split("|")]
    ok = True
    for c in cells:
        if not c:
            ok = False
            break
        # dashes with optional leading/trailing colon
        if not re.fullmatch(r":?-{3,}:?", c):
            ok = False
            break
    return ok


def parse_pipe_row(line: str) -> List[str]:
    s = line.strip().strip("\n").strip()
    # allow rows without leading/trailing pipe
    if "|" not in s:
        return [s]
    s = s.strip("|")
    cells = [c.strip() for c in s.split("|")]
    return cells


def try_parse_table(lines: List[str], i: int) -> Tuple[Optional[Table], int]:
    """
    Best-effort parse a Markdown pipe table starting at lines[i].
    Returns (table, new_index). If no table, returns (None, i).
    """
    if i + 1 >= len(lines):
        return None, i

    header_line = lines[i].rstrip("\n")
    sep_line = lines[i + 1].rstrip("\n")
    if "|" not in header_line:
        return None, i
    if not is_table_separator(sep_line):
        return None, i

    header = parse_pipe_row(header_line)
    rows: List[List[str]] = []
    j = i + 2
    while j < len(lines):
        line = lines[j].rstrip("\n")
        if not line.strip():
            break
        if "|" not in line:
            break
        rows.append(parse_pipe_row(line))
        j += 1

    # normalize row widths
    width = max(len(header), *(len(r) for r in rows)) if rows else len(header)
    header += [""] * (width - len(header))
    for r in rows:
        r += [""] * (width - len(r))

    return Table(header=header, rows=rows), j


def set_document_defaults(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 11) -> None:
    """
    Set basic default font for the document (English-friendly).
    """
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)

    # Also set East Asian font to keep Word happy if any mixed chars exist
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_blockquote_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Pt(18)  # subtle indent
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_list_item(doc: Document, text: str, ordered: bool, level: int) -> None:
    # Use built-in list styles where possible. Word list styling varies by template.
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(text, style=style)
    if level > 0:
        p.paragraph_format.left_indent = Pt(18 * level)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_table(doc: Document, table: Table) -> None:
    cols = len(table.header)
    t = doc.add_table(rows=1, cols=cols)
    t.style = "Table Grid"

    hdr_cells = t.rows[0].cells
    for c, val in enumerate(table.header):
        hdr_cells[c].text = val

    for row in table.rows:
        row_cells = t.add_row().cells
        for c in range(cols):
            row_cells[c].text = row[c] if c < len(row) else ""


def md_to_docx(md_text: str, output_path: str, strip_trace: bool = True, strip_all_comments: bool = False,
               title: Optional[str] = None) -> None:
    doc = Document()
    set_document_defaults(doc)

    if title:
        doc.add_heading(title, level=0)

    if strip_trace or strip_all_comments:
        md_text = strip_trace_blocks(md_text, strip_all_comments=strip_all_comments)

    lines = md_text.splitlines(keepends=True)

    in_code = False
    code_fence = None

    buf_paragraph: List[str] = []

    def flush_paragraph_buffer():
        nonlocal buf_paragraph
        text = "".join(buf_paragraph).strip()
        if text:
            # collapse internal newlines as spaces for Word paragraphs
            text = re.sub(r"\s*\n\s*", " ", text).strip()
            doc.add_paragraph(text)
        buf_paragraph = []

    i = 0
    while i < len(lines):
        line_raw = lines[i]
        line = line_raw.rstrip("\n")

        # code fences
        fence_match = re.match(r"^\s*```(.*)\s*$", line)
        if fence_match:
            flush_paragraph_buffer()
            if not in_code:
                in_code = True
                code_fence = "```"
            else:
                in_code = False
                code_fence = None
            i += 1
            continue

        if in_code:
            add_code_paragraph(doc, line.rstrip("\r"))
            i += 1
            continue

        # blank line -> paragraph break
        if not line.strip():
            flush_paragraph_buffer()
            i += 1
            continue

        # tables (best-effort)
        table, new_i = try_parse_table(lines, i)
        if table is not None:
            flush_paragraph_buffer()
            add_table(doc, table)
            i = new_i
            continue

        # headings
        h = re.match(r"^(#{1,6})\s+(.*)\s*$", line)
        if h:
            flush_paragraph_buffer()
            level = len(h.group(1))
            text = h.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # blockquote
        bq = re.match(r"^\s*>\s?(.*)$", line)
        if bq:
            flush_paragraph_buffer()
            add_blockquote_paragraph(doc, bq.group(1).strip())
            i += 1
            continue

        # ordered list
        ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if ol:
            flush_paragraph_buffer()
            indent = len(ol.group(1).expandtabs(4))
            level = indent // 2  # heuristic
            add_list_item(doc, ol.group(3).strip(), ordered=True, level=level)
            i += 1
            continue

        # unordered list
        ul = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if ul:
            flush_paragraph_buffer()
            indent = len(ul.group(1).expandtabs(4))
            level = indent // 2  # heuristic
            add_list_item(doc, ul.group(2).strip(), ordered=False, level=level)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            flush_paragraph_buffer()
            p = doc.add_paragraph("—" * 20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # default: accumulate into a paragraph buffer
        buf_paragraph.append(line_raw)
        i += 1

    flush_paragraph_buffer()
    doc.save(output_path)


def main():
    ap = argparse.ArgumentParser(description="Convert Markdown to DOCX (Word) with simple rules.")
    ap.add_argument("input_md", help="Input Markdown file path")
    ap.add_argument("-o", "--output", required=False, help="Output .docx path (default: input basename .docx)")
    ap.add_argument("--title", default=None, help="Optional document title to add at top (Heading 0)")
    ap.add_argument("--keep-trace", action="store_true", help="Keep <!-- TRACE ... --> blocks (do not strip)")
    ap.add_argument("--strip-all-comments", action="store_true", help="Strip all HTML comments (not only TRACE)")
    args = ap.parse_args()

    input_path = args.input_md
    output_path = args.output
    if not output_path:
        output_path = re.sub(r"\.md$", "", input_path, flags=re.IGNORECASE) + ".docx"

    with open(input_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    md_to_docx(
        md_text=md_text,
        output_path=output_path,
        strip_trace=(not args.keep_trace),
        strip_all_comments=args.strip_all_comments,
        title=args.title,
    )

    print(f"[OK] Wrote: {output_path}")


if __name__ == "__main__":
    main()