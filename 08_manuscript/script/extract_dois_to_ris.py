#!/usr/bin/env python3
# extract_dois_to_ris.py
# -*- coding: utf-8 -*-

import re
import sys
import time
from collections import OrderedDict
from pathlib import Path

import requests
from docx import Document
from tqdm import tqdm

# ------------ 设置：当前目录为目标 ------------
FOLDER = Path("/Volumes/Samsung_SSD_990_PRO_2TB_Media/EphB1/08_manuscript")
OUT_RIS = FOLDER / "references_out.ris"
OUT_SKIP = FOLDER / "skipped_files.txt"
# More tolerant DOI extraction:
# Matches {10.xxxx/...} and {doi:10.xxxx/...} (case-insensitive)
DOI_CORE = r"10\.\d{4,9}/[-._;()/:A-Z0-9]+"
REGEX_DOI = re.compile(r"\{\s*(?:doi\s*:\s*)?(" + DOI_CORE + r")\s*\}", flags=re.IGNORECASE)
REGEX_DOI_BARE = re.compile(r"(?<!\{)(?<!\w)(" + DOI_CORE + r")(?!\w)", flags=re.IGNORECASE)

HEADERS = {
    "User-Agent": "doi-ris-extractor/0.2 (mailto:your_email@example.com)",
    "Accept": "application/x-research-info-systems",
}
# ---------------------------------------------

def extract_dois_from_doc(doc_path: Path, skipped_list: list[str]) -> list[str]:
    try:
        doc = Document(doc_path)

        def iter_text():
            # Main body paragraphs
            for p in doc.paragraphs:
                if p.text:
                    yield p.text

            # Tables (all cells)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text:
                                yield p.text

            # Headers / footers
            for section in doc.sections:
                hdr = section.header
                ftr = section.footer
                for p in getattr(hdr, "paragraphs", []):
                    if p.text:
                        yield p.text
                for tbl in getattr(hdr, "tables", []):
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text:
                                    yield p.text
                for p in getattr(ftr, "paragraphs", []):
                    if p.text:
                        yield p.text
                for tbl in getattr(ftr, "tables", []):
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text:
                                    yield p.text

        found: list[str] = []
        for chunk in iter_text():
            found.extend(REGEX_DOI.findall(chunk))
            found.extend(REGEX_DOI_BARE.findall(chunk))

        return found
    except Exception as e:
        skipped_list.append(f"{doc_path.name} - {e}")
        return []

def fetch_ris(doi: str) -> str:
    url = f"https://api.crossref.org/works/{doi}/transform/application/x-research-info-systems"
    try:
        r = requests.get(url, headers=HEADERS, timeout=20)
        if r.status_code == 200 and "TY  -" in r.text:
            return r.text.strip()
    except Exception:
        pass
    return f"TY  - GEN\nDO  - {doi}\nER  -"

def main():
    # Prefer command-line provided docx paths; fallback to scanning FOLDER
    if len(sys.argv) > 1:
        docx_files = [Path(p) for p in sys.argv[1:]]
    else:
        docx_files = sorted(FOLDER.rglob("*.docx"))

    OUT_RIS.parent.mkdir(parents=True, exist_ok=True)

    if not docx_files:
        print("❌ 当前目录没有找到 .docx 文件")
        sys.exit(1)

    all_dois = OrderedDict()
    skipped = []

    print(f"🔍 正在扫描 {len(docx_files)} 个 Word 文件 …")
    for docx in docx_files:
        for doi in extract_dois_from_doc(docx, skipped):
            norm = doi.strip().strip(". ;,)]}>\"")
            all_dois[norm.lower()] = None  # 去重 + 保顺序

    if not all_dois:
        print("⚠ 没有发现任何 {DOI} 引用")
    else:
        print(f"🚀 开始抓取 Crossref，共 {len(all_dois)} 个唯一 DOI …")
        for doi in tqdm(all_dois, unit="DOI"):
            all_dois[doi] = fetch_ris(doi)
            time.sleep(0.1)  # 礼貌性等待

        OUT_RIS.write_text("\n\n".join(all_dois.values()), encoding="utf-8")
        print(f"✅ 已生成 RIS 文件：{OUT_RIS}")

    if skipped:
        OUT_SKIP.write_text("\n".join(skipped), encoding="utf-8")
        print(f"⚠ 以下文件无法解析，已写入：{OUT_SKIP}")
    else:
        print("✅ 所有 Word 文件均成功解析。")

if __name__ == "__main__":
    main()
