#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
auto_doi_from_ris.py  —— 升级版
1. 把正文 [序号] 替换成 {doi}{doi}…（每个 DOI 独立包裹）
2. References 末尾仍在首段追加 {doi}
3. 文档末尾追加 “Full DOI Reference List”：序号 + 标题 + {doi}
"""
import sys, re
import rispy
from docx import Document

REF_HEADINGS = {"references", "reference", "参考文献"}

# ---------- RIS → 映射 ----------
def load_ris(path):
    with open(path, 'r', encoding='utf-8') as f:
        entries = rispy.load(f)
    num2doi = {}
    for i, rec in enumerate(entries, 1):
        doi = rec.get('doi')
        if not doi:
            url = rec.get('url', '')
            doi = url[url.find('10.'):] if '10.' in url else f'NO_DOI_{i}'
        num2doi[i] = doi
    return num2doi, entries

# ---------- Word 处理 ----------
def locate_references(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() in REF_HEADINGS:
            return i
    raise ValueError("找不到 References 标题")

def append_doi_to_refs(doc, ref_start, num2doi):
    idx = 1
    for p in doc.paragraphs[ref_start+1:]:
        if not p.text.strip():
            continue
        if idx > len(num2doi):
            break
        if f"{{{num2doi[idx]}}}" not in p.text:
            p.text = p.text.rstrip() + f" {{{num2doi[idx]}}}"
        idx += 1

def replace_in_body(doc, ref_start, num2doi):
    cite_pat = re.compile(r'\[([0-9,\s\u00A0\-–—]+)\]')

    def repl(m):
        inside = re.sub(r'[\s\u00A0]', '', m.group(1))
        parts  = inside.split(',')
        idxs   = []
        for part in parts:
            if not part: continue
            if re.search(r'[-–—]', part):
                a, b = re.split(r'[-–—]', part)
                idxs.extend(range(int(a), int(b)+1))
            else:
                idxs.append(int(part))
        # —— 每 DOI 独立 { } —— #
        return "".join(f'{{{num2doi.get(i, f"NO_DOI_{i}")}}}' for i in idxs)

    for p in doc.paragraphs[:ref_start]:
        for run in p.runs:
            run.text = cite_pat.sub(repl, run.text)

def append_master_list(doc, num2doi, ris_entries):
    doc.add_page_break()
    h = doc.add_paragraph('Full DOI Reference List')
    h.style = 'Heading 1'

    for i, entry in enumerate(ris_entries, 1):
        title = entry.get('title', [''])[0] if isinstance(entry.get('title'), list) else entry.get('title', '')
        doc.add_paragraph(f"{i}. {title} {{{num2doi.get(i, f'NO_DOI_{i}')}}}")

# ---------- main ----------
def main(src_doc, ris_file, dst_doc):
    num2doi, ris_entries = load_ris(ris_file)
    doc = Document(src_doc)
    ref_idx = locate_references(doc)

    append_doi_to_refs(doc, ref_idx, num2doi)
    replace_in_body(doc, ref_idx, num2doi)
    append_master_list(doc, num2doi, ris_entries)

    doc.save(dst_doc)
    print("✅ 处理完成:", dst_doc)

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print("用法: python3 auto_doi_from_ris.py 原文.docx refs.ris 输出.docx")
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3])
